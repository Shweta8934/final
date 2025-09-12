import streamlit as st
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
import docx
import pandas as pd
import json
import io
import zipfile
import tempfile
import os
import requests
from typing import Tuple, Optional

class FileProcessor:
    """Handle various file types and extract readable content"""
    
    def __init__(self):
        self.supported_types = [
            "png", "jpg", "jpeg", "gif", "bmp", "tiff",  # Images
            "pdf",  # PDF
            "docx", "doc",  # Word documents
            "txt", "md",  # Text files
            "csv", "xlsx", "xls",  # Spreadsheets
            "json",  # JSON files
            "zip"  # ZIP files (limited support)
        ]
    
    def extract_text_from_image(self, image_file) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(image_file)
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            text = pytesseract.image_to_string(image, lang='eng')
            return text.strip()
        except Exception as e:
            raise Exception(f"Error processing image: {str(e)}")
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text from PDF using OCR"""
        try:
            # Convert PDF pages to images
            pages = convert_from_bytes(pdf_file.read())
            full_text = ""
            
            for i, page in enumerate(pages):
                try:
                    page_text = pytesseract.image_to_string(page, lang='eng')
                    full_text += f"\n--- Page {i+1} ---\n{page_text}\n"
                except Exception as e:
                    full_text += f"\n--- Page {i+1} (Error: {str(e)}) ---\n"
            
            return full_text.strip()
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(docx_file)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            return "\n".join(full_text)
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def extract_text_from_txt(self, txt_file) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    txt_file.seek(0)  # Reset file pointer
                    content = txt_file.read().decode(encoding)
                    return content.strip()
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with supported encodings")
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    def extract_text_from_csv(self, csv_file) -> str:
        """Extract and format data from CSV file"""
        try:
            df = pd.read_csv(csv_file)
            
            # Create a readable summary of the CSV
            summary = f"CSV File Summary:\n"
            summary += f"- Rows: {len(df)}\n"
            summary += f"- Columns: {len(df.columns)}\n"
            summary += f"- Column Names: {', '.join(df.columns.tolist())}\n\n"
            
            # Show first few rows
            summary += "First 5 rows:\n"
            summary += df.head().to_string(index=False)
            
            # Show basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                summary += "\n\nBasic Statistics for Numeric Columns:\n"
                summary += df[numeric_cols].describe().to_string()
            
            return summary
        except Exception as e:
            raise Exception(f"Error processing CSV file: {str(e)}")
    
    def extract_text_from_excel(self, excel_file) -> str:
        """Extract and format data from Excel file"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            full_text = f"Excel File Summary:\n"
            full_text += f"- Number of sheets: {len(excel_data)}\n"
            
            for sheet_name, df in excel_data.items():
                full_text += f"\n--- Sheet: {sheet_name} ---\n"
                full_text += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                full_text += f"Column Names: {', '.join(df.columns.tolist())}\n"
                
                # Show first few rows
                full_text += "\nFirst 5 rows:\n"
                full_text += df.head().to_string(index=False)
                full_text += "\n"
            
            return full_text
        except Exception as e:
            raise Exception(f"Error processing Excel file: {str(e)}")
    
    def extract_text_from_json(self, json_file) -> str:
        """Extract and format data from JSON file"""
        try:
            data = json.load(json_file)
            
            # Format JSON data in a readable way
            formatted_json = json.dumps(data, indent=2, ensure_ascii=False)
            
            summary = f"JSON File Content:\n"
            summary += f"Data type: {type(data).__name__}\n"
            
            if isinstance(data, dict):
                summary += f"Number of keys: {len(data)}\n"
                summary += f"Keys: {', '.join(list(data.keys())[:10])}{'...' if len(data) > 10 else ''}\n\n"
            elif isinstance(data, list):
                summary += f"Number of items: {len(data)}\n\n"
            
            summary += "Formatted Content:\n"
            summary += formatted_json[:2000]  # Limit length
            
            if len(formatted_json) > 2000:
                summary += "\n... (content truncated for display)"
            
            return summary
        except Exception as e:
            raise Exception(f"Error processing JSON file: {str(e)}")
    
    def process_file(self, uploaded_file) -> Tuple[str, str]:
        """
        Process uploaded file and extract text content
        Returns: (extracted_text, file_info)
        """
        if not uploaded_file:
            return "", ""
        
        file_name = uploaded_file.name
        file_type = file_name.split('.')[-1].lower()
        file_size = uploaded_file.size
        
        # File info for display
        file_info = f"📁 **File:** {file_name}\n📏 **Size:** {file_size:,} bytes\n📄 **Type:** {file_type.upper()}"
        
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            if file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                extracted_text = self.extract_text_from_image(uploaded_file)
                
            elif file_type == 'pdf':
                extracted_text = self.extract_text_from_pdf(uploaded_file)
                
            elif file_type == 'docx':
                extracted_text = self.extract_text_from_docx(uploaded_file)
                
            elif file_type in ['txt', 'md']:
                extracted_text = self.extract_text_from_txt(uploaded_file)
                
            elif file_type == 'csv':
                extracted_text = self.extract_text_from_csv(uploaded_file)
                
            elif file_type in ['xlsx', 'xls']:
                extracted_text = self.extract_text_from_excel(uploaded_file)
                
            elif file_type == 'json':
                extracted_text = self.extract_text_from_json(uploaded_file)
                
            else:
                raise Exception(f"Unsupported file type: {file_type}")
            
            if not extracted_text or extracted_text.isspace():
                raise Exception("No readable content found in the file")
            
            return extracted_text, file_info
            
        except Exception as e:
            error_msg = f"❌ Error processing file: {str(e)}"
            return "", error_msg

def render_file_upload_section(subject: str) -> Tuple[str, str, str]:
    """
    Render file upload UI and return extracted content
    Returns: (extracted_text, file_info, question)
    """
    processor = FileProcessor()
    
    # Create two columns for layout
    col1, col2 = st.columns([1, 1])
    
    with col1:
        # st.markdown("### 📎 Upload File (Optional)")
        # st.info(f"💡 Upload files related to **{subject}**")
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Choose a file to analyze",
            type=processor.supported_types,
            help="Images, PDF, Word, Text, CSV, Excel, JSON",
            key="file_uploader"
        )
    
    with col2:
        # st.markdown("### ❓ Ask Your Question")
        # st.info("💭 Type your question (with or without file)")
        
        # Question input - ALWAYS VISIBLE
        question = st.text_area(
            "Type your question here",
            height=120,
            placeholder=f"Example: How do I solve 2x + 5 = 15?" if subject == "Math" 
            else f"Example: What is photosynthesis?" if subject == "Science"
            else f"Example: What is a metaphor?" if subject == "English"
            else "Ask anything you'd like to learn about!",
            key="question_input"
        )
    
    # Process uploaded file
    extracted_text = ""
    file_info = ""
    
    if uploaded_file:
        with st.spinner("🔄 Processing your file..."):
            extracted_text, file_info = processor.process_file(uploaded_file)
        
        if extracted_text:
            st.success("✅ File processed successfully!")
            
            # Show file info in a compact way
            st.markdown(f"**📁 File:** {uploaded_file.name} ({uploaded_file.size:,} bytes)")
            
            # Show extracted content - ALWAYS EXPANDED for visibility
            st.markdown("### 📄 Extracted Content from File:")
            
            if len(extracted_text) > 500:
                # Show preview + expandable full content
                st.text_area(
                    "Content Preview (first 500 chars)",
                    value=extracted_text[:500] + "...",
                    height=150,
                    disabled=True,
                    key="content_preview"
                )
                
                with st.expander("📋 View Full Extracted Content", expanded=False):
                    st.text_area(
                        "Full Content",
                        value=extracted_text,
                        height=300,
                        disabled=True,
                        key="full_content"
                    )
                st.info(f"📊 Total content length: {len(extracted_text):,} characters")
            else:
                st.text_area(
                    "Extracted Content",
                    value=extracted_text,
                    height=min(200, len(extracted_text.split('\n')) * 25),
                    disabled=True,
                    key="extracted_content"
                )
            
        else:
            st.error(file_info)  # Show error message
    
    return extracted_text, file_info, question

def get_file_analysis_prompt(extracted_text: str, question: str, subject: str, grade: str) -> str:
    """
    Generate specialized prompt for file-based questions
    """
    if extracted_text and question:
        return f"""
As an AI tutor for {grade} students studying {subject}, I need to help analyze uploaded content and answer specific questions about it.

UPLOADED CONTENT:
{extracted_text}

STUDENT'S QUESTION:
{question}

Please provide a comprehensive educational response that:
1. Analyzes the uploaded content in relation to the question
2. Explains relevant concepts clearly for a {grade} student
3. Provides step-by-step explanations where needed
4. Relates the content to {subject} curriculum
5. Offers additional learning suggestions
6. Encourages further exploration of the topic

Format your response with clear sections and use emojis for better organization.
"""
    elif extracted_text:
        return f"""
As an AI tutor for {grade} students studying {subject}, I need to analyze uploaded content and provide educational insights.

UPLOADED CONTENT:
{extracted_text}

Please provide a comprehensive educational analysis that:
1. Summarizes the key points from the uploaded content
2. Explains the main concepts for a {grade} student level
3. Identifies important learning objectives
4. Relates the content to {subject} curriculum
5. Suggests questions the student might want to explore
6. Provides additional learning resources or activities

Format your response with clear sections and use emojis for better organization.
"""
    else:
        return f"""
As an AI tutor for {grade} students studying {subject}, please answer the following question:

{question}

Provide a comprehensive educational response appropriate for {grade} level.
"""

# Example usage function for testing
def test_file_processor():
    """Test function to demonstrate usage"""
    processor = FileProcessor()
    print("Supported file types:", processor.supported_types)
    return processor